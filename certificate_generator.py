#!/usr/bin/env python3
"""Generate filled certificates from Word and PDF templates."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import fitz
from docx import Document

try:
    import arabic_reshaper
    from bidi.algorithm import get_display

    def _shape_arabic(text: str) -> str:
        return get_display(arabic_reshaper.reshape(text))
except ImportError:
    def _shape_arabic(text: str) -> str:
        return text


def _contains_arabic(text: str) -> bool:
    return bool(re.search(r"[\u0600-\u06FF]", text))


def _prepare_pdf_text(text: str) -> str:
    if _contains_arabic(text):
        return _shape_arabic(text)
    return text


def _resolve_base_dir() -> Path:
    """App root: source tree, or PyInstaller bundle / exe folder."""
    import sys

    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        meipass = Path(getattr(sys, "_MEIPASS", exe_dir))
        for candidate in (meipass, exe_dir / "_internal", exe_dir):
            if (candidate / "certificate_ui.py").exists() or (
                candidate / "certificates templates"
            ).exists():
                return candidate
        return meipass
    return Path(__file__).resolve().parent


def _resolve_writable_dir() -> Path:
    """Writable folder for outputs (next to the .exe when frozen)."""
    import sys

    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


BASE_DIR = _resolve_base_dir()
WRITE_DIR = _resolve_writable_dir()
TEMPLATES_DIR = BASE_DIR / "certificates templates"
OUTPUT_DIR = WRITE_DIR / "generated_certificates"
BULK_DIR = WRITE_DIR / "bulk_templates"
PREVIEW_CACHE_DIR = WRITE_DIR / "preview_cache"
STATIC_PREVIEWS_DIR = BASE_DIR / "assets" / "previews"
HONORARY_PORTRAIT_SOURCE = "archive/نموذج شهادة 2.pdf"
HONORARY_LANDSCAPE_SOURCE = "شهادة شرفية.pdf"


def _default_arabic_fonts() -> tuple[str, str]:
    import os

    candidates = [
        BASE_DIR / "fonts" / "Amiri-Regular.ttf",
        BASE_DIR / "fonts" / "ScheherazadeNew-Regular.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    if os.name == "nt":
        windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
        candidates.extend(
            [
                windir / "Fonts" / "arabtype.ttf",
                windir / "Fonts" / "tradbdo.ttf",
                windir / "Fonts" / "arial.ttf",
            ]
        )
    regular = next((str(p) for p in candidates if p.exists()), "DejaVuSans.ttf")
    bold_candidates = [
        BASE_DIR / "fonts" / "Amiri-Bold.ttf",
        BASE_DIR / "fonts" / "ScheherazadeNew-Bold.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    bold = next((str(p) for p in bold_candidates if p.exists()), regular)
    return regular, bold


ARABIC_FONT, ARABIC_FONT_BOLD = _default_arabic_fonts()
PREFERRED_ARABIC_FONTS = (
    "Amiri",
    "Scheherazade New",
    "Arabic Typesetting",
    "Noto Naskh Arabic",
    "DejaVu Sans",
)


@dataclass
class CertificateField:
    id: str
    label_ar: str
    label_en: str
    label_fr: str = ""
    field_type: str = "text"
    default: str = ""
    required: bool = True


@dataclass
class CertificateTemplate:
    id: str
    name_ar: str
    name_en: str
    name_fr: str
    source_file: str
    output_format: str
    fields: list[CertificateField]
    category: str = "docx"


def _dots(n: int) -> str:
    return "." * n


TEMPLATES: dict[str, CertificateTemplate] = {
    "appreciation": CertificateTemplate(
        id="appreciation",
        name_ar="شهادة شكر و تقدير",
        name_en="Certificate of Appreciation",
        name_fr="Certificat de remerciement",
        source_file="شهادة شكر و تقدير.docx",
        output_format="docx",
        fields=[
            CertificateField("recipient_name", "اسم المستفيد", "Recipient name", "Nom du bénéficiaire"),
            CertificateField("training_field", "مجال الدورة التكوينية", "Training field", "Domaine de formation"),
            CertificateField("date_from", "تاريخ البداية", "Start date", "Date de début"),
            CertificateField("date_to", "تاريخ النهاية", "End date", "Date de fin"),
            CertificateField("coordinating_orgs", "الجهات المتنسقة", "Coordinating organizations", "Organismes coordinateurs"),
        ],
    ),
    "condolence_individual": CertificateTemplate(
        id="condolence_individual",
        name_ar="تعزية (للشخص)",
        name_en="Condolence (Individual)",
        name_fr="Condoléances",
        source_file="تعزية .docx",
        output_format="docx",
        category="condolence",
        fields=[
            CertificateField("recipient_name", "اسم المستفيد", "Recipient name", "Nom du bénéficiaire"),
            CertificateField("rank_workplace", "الرتبة و مكان العمل", "Rank & workplace", "Grade et lieu de travail"),
            CertificateField("deceased_name", "اسم المتوفى(ة)", "Deceased name", "Nom du défunt"),
        ],
    ),
    "congratulations": CertificateTemplate(
        id="congratulations",
        name_ar="تهنئة",
        name_en="Congratulations Letter",
        name_fr="Lettre de félicitations",
        source_file="تهنئة.docx",
        output_format="docx",
        fields=[
            CertificateField("recipient_title", "الاسم الكامل | اللقب", "full name | title", "Nom complet | Titre"),
            CertificateField("council_location", "موقع المجلس", "council location", "Lieu du conseil"),
        ],
    ),
    "honorary": CertificateTemplate(
        id="honorary",
        name_ar="شهادة شرفية",
        name_en="Honorary Certificate",
        name_fr="Certificat honorifique",
        source_file=HONORARY_LANDSCAPE_SOURCE,
        output_format="pdf",
        fields=[
            CertificateField("recipient_name", "اسم المستفيد", "Recipient full name", "Nom complet"),
            CertificateField("title", "اللقب (السيد/السيدة)", "Title (Mr/Ms)", "Civilité", default="السيد(ة)", required=False),
            CertificateField("rank", "الرتبة", "Rank", "Grade"),
            CertificateField("job", "الوظيفة", "Job position", "Poste"),
            CertificateField("issue_date", "تاريخ الإصدار", "Issue date", "Date", default=date.today().strftime("%d/%m/%Y")),
        ],
    ),
}


def list_templates() -> list[CertificateTemplate]:
    return list(TEMPLATES.values())


def get_template(template_id: str) -> CertificateTemplate:
    if template_id not in TEMPLATES:
        raise KeyError(f"Unknown template: {template_id}")
    return TEMPLATES[template_id]


def _replace_dots(text: str, replacements: list[tuple[str, int]]) -> str:
    """Replace dot sequences in order with provided values."""
    result = text
    for value, count in replacements:
        pattern = re.escape(_dots(count))
        result, count_done = re.subn(pattern, value, result, count=1)
        if count_done == 0:
            pattern = r"\.{5,}"
            result, count_done = re.subn(pattern, value, result, count=1)
    return result


def _replace_dots_in_runs(paragraph, value: str) -> bool:
    """Replace the first dotted placeholder inside existing runs (keeps font size/style)."""
    for run in paragraph.runs:
        text = run.text or ""
        if re.search(r"\.{5,}", text):
            # Keep a space before the value when the run is only dots (label is in a previous run).
            if re.fullmatch(r"\s*\.{5,}\s*", text):
                run.text = f" {value} "
            else:
                run.text = re.sub(r"\.{5,}\s*", value + " ", text, count=1)
            return True
    return False


def _set_paragraph_text_keep_style(paragraph, new_text: str) -> None:
    """Overwrite paragraph text while keeping the first run's formatting."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _apply_docx_rules(
    doc: Document,
    rules: dict[int, list[tuple[str, int]] | str],
    data: dict[str, str],
) -> None:
    for index, rule in rules.items():
        if index >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[index]
        if isinstance(rule, str):
            # Prefer in-place dot replacement when the rule is only filling blanks.
            field_ids = re.findall(r"\{(\w+)\}", rule)
            if field_ids and _replace_dots_in_runs(paragraph, data[field_ids[0]]) and len(field_ids) == 1:
                # If the template already has the label text, only the dots were replaced.
                # When the formatted rule differs from the template label, fall back below.
                expected_prefix = rule.split("{")[0]
                if paragraph.text.startswith(expected_prefix.strip()[:8]) or expected_prefix.strip() in paragraph.text:
                    continue
                _set_paragraph_text_keep_style(paragraph, rule.format(**data))
            else:
                _set_paragraph_text_keep_style(paragraph, rule.format(**data))
            continue
        for field_id, _dot_count in rule:
            if not _replace_dots_in_runs(paragraph, data[field_id]):
                # Fallback for odd layouts
                paragraph.text = _replace_dots(paragraph.text, [(data[field_id], _dot_count)])
            break


def _trim_condolence_doc(doc: Document, keep: str) -> None:
    """Keep only one condolence letter variant in the document."""
    if keep == "individual":
        for _ in range(len(doc.paragraphs) - 1, 14, -1):
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)
    elif keep == "family":
        for _ in range(13, -1, -1):
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)


DOCX_RULES: dict[str, dict[int, list[tuple[str, int]] | str]] = {
    "condolence_individual": {
        8: [("recipient_name", 21)],
        9: [("rank_workplace", 33)],
        12: [("deceased_name", 10)],
    },
    "condolence_family": {
        28: [("family_name", 21)],
        29: [("rank_workplace", 33)],
        32: [("deceased_name", 10)],
    },
}


def _resolve_data(template: CertificateTemplate, data: dict[str, Any]) -> dict[str, str]:
    resolved: dict[str, str] = {}
    for field_def in template.fields:
        value = data.get(field_def.id, field_def.default)
        if field_def.required and not str(value).strip():
            raise ValueError(f"Missing required field: {field_def.label_ar}")
        resolved[field_def.id] = str(value).strip()
    return resolved


def _derive_year(data: dict[str, str]) -> str:
    """Pull a year from date fields when the Année UI field is omitted."""
    for key in ("date_to", "date_from"):
        match = re.search(r"(20\d{2}|19\d{2})", data.get(key, ""))
        if match:
            return match.group(1)
    return str(date.today().year)


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^\w\s\-]", "", value, flags=re.UNICODE)
    cleaned = re.sub(r"\s+", "_", cleaned.strip())
    return cleaned[:80] or "certificate"


def _set_run_half_points(run, half_points: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    rPr = run._element.get_or_add_rPr()
    for tag in ("sz", "szCs"):
        node = rPr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            rPr.append(node)
        node.set(qn("w:val"), str(half_points))
    run.font.size = Pt(half_points / 2)


def _bump_runs_at_most(paragraph, max_half_points: int, new_half_points: int) -> None:
    from docx.oxml.ns import qn

    for run in paragraph.runs:
        rPr = run._element.rPr
        current = None
        if rPr is not None:
            for tag in ("sz", "szCs"):
                node = rPr.find(qn(f"w:{tag}"))
                if node is not None:
                    try:
                        current = int(node.get(qn("w:val")))
                        break
                    except (TypeError, ValueError):
                        pass
        if current is None or current <= max_half_points:
            _set_run_half_points(run, new_half_points)


def _iter_paragraphs(doc: Document):
    """Yield every paragraph, including those nested in tables."""
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _enlarge_appreciation_text(doc: Document) -> None:
    """Slightly enlarge body/name only — keep single-page layout."""
    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("شهادة") or text.startswith("◆") or "ختم" in text or text == "التوقيع":
            continue
        # Template body ≈ 32 half-points (16 pt) → 34 (17 pt), a small bump only.
        _bump_runs_at_most(paragraph, max_half_points=32, new_half_points=34)


def _distribute_appreciation_layout(doc: Document) -> None:
    """
    Push title-and-below content downward so the landscape page is balanced,
    without overflowing onto a second page.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    def _spacing(paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        node = pPr.find(qn("w:spacing"))
        if node is None:
            node = OxmlElement("w:spacing")
            pPr.append(node)
        return node

    def _make_spacer(before: int = 120, after: int = 120):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(before))
        sp.set(qn("w:after"), str(after))
        pPr.append(sp)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        p.append(pPr)
        return p

    # Keep at most one blank spacer between blocks (template has several).
    blanks: list = []
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip():
            blanks.append(paragraph)
            continue
        if len(blanks) > 1:
            for extra in blanks[1:]:
                el = extra._element
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
        blanks = []
    if len(blanks) > 1:
        for extra in blanks[1:]:
            el = extra._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    # Insert spacers just before the title to shift the whole lower block down.
    title_el = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("شهادة"):
            title_el = paragraph._element
            break
    if title_el is not None and title_el.getparent() is not None:
        parent = title_el.getparent()
        for _ in range(3):
            parent.insert(parent.index(title_el), _make_spacer(120, 120))

    title_seen = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        sp = _spacing(paragraph)

        if text.startswith("شهادة"):
            title_seen = True
            sp.set(qn("w:before"), "160")
            sp.set(qn("w:after"), "120")
            sp.set(qn("w:line"), "260")
            sp.set(qn("w:lineRule"), "auto")
            continue

        if not title_seen:
            continue

        if text.startswith("◆"):
            sp.set(qn("w:before"), "100")
            sp.set(qn("w:after"), "120")
        elif text.startswith("السيد"):
            sp.set(qn("w:before"), "140")
            sp.set(qn("w:after"), "140")
            sp.set(qn("w:line"), "280")
            sp.set(qn("w:lineRule"), "auto")
        elif text.startswith("تقديرا") or text.startswith("المنجزة"):
            sp.set(qn("w:before"), "90")
            sp.set(qn("w:after"), "90")
            sp.set(qn("w:line"), "280")
            sp.set(qn("w:lineRule"), "auto")
        elif "ختم" in text:
            sp.set(qn("w:before"), "160")
            sp.set(qn("w:after"), "100")
        elif text == "التوقيع":
            sp.set(qn("w:before"), "80")
            sp.set(qn("w:after"), "40")
        elif not text:
            sp.set(qn("w:before"), "100")
            sp.set(qn("w:after"), "100")

    section = doc.sections[0]
    section.bottom_margin = Cm(0.55)
    section.top_margin = Cm(0.7)


def _enlarge_congratulations_fields(doc: Document, recipient_title: str, council_location: str) -> None:
    """Make Nom complet | Titre and Lieu du conseil larger than body text."""
    targets = {value.strip() for value in (recipient_title, council_location) if value and value.strip()}
    if not targets:
        return
    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text
        if any(target in text for target in targets):
            for run in paragraph.runs:
                _set_run_half_points(run, 44)


def _docx_uses_jinja(source: Path) -> bool:
    import zipfile

    with zipfile.ZipFile(source) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return "{{" in xml and "}}" in xml


def _font_available(family: str) -> bool:
    """True if the font can be used for DOCX → PDF (bundled, system, or fc-list)."""
    import os

    bundled = {
        "Amiri": BASE_DIR / "fonts" / "Amiri-Regular.ttf",
        "Scheherazade New": BASE_DIR / "fonts" / "ScheherazadeNew-Regular.ttf",
        "DejaVu Sans": Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    }
    bundled_path = bundled.get(family)
    if bundled_path is not None and bundled_path.exists():
        return True

    if os.name == "nt":
        windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
        fonts_dir = windir / "Fonts"
        aliases = {
            "Amiri": ("amiri-regular.ttf", "amiri.ttf"),
            "Scheherazade New": ("scheherazadenew-regular.ttf", "scheherazade new.ttf"),
            "Arabic Typesetting": ("arabtype.ttf",),
            "Noto Naskh Arabic": ("notonaskharabic-regular.ttf", "NotoNaskhArabic-Regular.ttf"),
            "DejaVu Sans": ("dejavusans.ttf",),
        }
        for name in aliases.get(family, ()):
            if (fonts_dir / name).exists():
                return True
        # Case-insensitive scan for family token.
        token = family.lower().replace(" ", "")
        try:
            for entry in fonts_dir.iterdir():
                if token in entry.name.lower().replace(" ", ""):
                    return True
        except OSError:
            pass

    import subprocess

    try:
        result = subprocess.run(
            ["fc-list", family, "file"],
            capture_output=True,
            text=True,
            check=False,
        )
        return bool(result.stdout.strip())
    except OSError:
        return False


def _best_arabic_font() -> str:
    for family in PREFERRED_ARABIC_FONTS:
        if _font_available(family):
            return family
    return "Amiri"


def _contains_arabic(text: str) -> bool:
    return any("\u0600" <= ch <= "\u06FF" for ch in text)


def _strip_run_italic(run) -> None:
    """Remove italic flags — Arial Italic has no Arabic glyphs in LibreOffice PDF."""
    from docx.oxml.ns import qn

    run.font.italic = False
    rPr = run._element.get_or_add_rPr()
    for tag in ("i", "iCs"):
        node = rPr.find(qn(f"w:{tag}"))
        if node is not None:
            rPr.remove(node)


def _run_is_italic(run) -> bool:
    from docx.oxml.ns import qn

    if run.font.italic:
        return True
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return False
    return rPr.find(qn("w:i")) is not None or rPr.find(qn("w:iCs")) is not None


def _fix_arabic_pdf_fonts(doc: Document, *, size_scale: float = 1.0) -> str:
    """
    Make Arabic runs PDF-safe for LibreOffice without blowing up layout.

    Only italic Arabic is remapped (Arial Italic → tofu □□□ in LO PDF).
    Non-italic Arial Arabic already renders fine and stays compact on one page.
    """
    font_name = _best_arabic_font()
    for paragraph in _iter_paragraphs(doc):
        for run in paragraph.runs:
            if not run.text or not _contains_arabic(run.text):
                continue
            if not _run_is_italic(run):
                continue
            _strip_run_italic(run)
            _set_run_font(run, font_name, size_scale=size_scale)
    return font_name


def _set_line_spacing(doc: Document, line: int = 276) -> None:
    """Set paragraph line spacing (240 = single, 276 ≈ 1.15, 288 = 1.2)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        pPr = paragraph._element.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def _set_run_font(run, font_name: str, size_scale: float = 1.0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), font_name)

    if size_scale == 1.0:
        return

    # Scale both ascii and complex-script sizes (Arabic uses w:szCs).
    for tag in ("sz", "szCs"):
        node = rPr.find(qn(f"w:{tag}"))
        if node is None:
            continue
        try:
            half_points = int(node.get(qn("w:val")))
        except (TypeError, ValueError):
            continue
        scaled = max(16, int(round(half_points * size_scale)))
        node.set(qn("w:val"), str(scaled))
        if tag == "sz" and run.font.size is not None:
            run.font.size = Pt(scaled / 2)


def _tighten_paragraph_spacing(doc: Document) -> None:
    """Reduce extra paragraph gaps so tall Arabic fonts still fit on one page."""
    from docx.oxml.ns import qn
    from docx.shared import Twips

    # Drop consecutive blank spacer paragraphs (keep at most one).
    blanks: list = []
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            blanks.append(paragraph)
        else:
            if len(blanks) > 1:
                for extra in blanks[1:]:
                    el = extra._element
                    parent = el.getparent()
                    if parent is not None:
                        parent.remove(el)
            blanks = []
    if len(blanks) > 1:
        for extra in blanks[1:]:
            el = extra._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    for paragraph in doc.paragraphs:
        pPr = paragraph._element.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            continue
        # Compact auto line spacing (240 = single).
        if spacing.get(qn("w:line")):
            try:
                line = int(spacing.get(qn("w:line")))
                spacing.set(qn("w:line"), str(max(200, min(line, 220))))
            except (TypeError, ValueError):
                pass
        for attr in ("before", "after"):
            key = qn(f"w:{attr}")
            if spacing.get(key):
                try:
                    val = int(spacing.get(key))
                    spacing.set(key, str(max(0, val // 2)))
                except (TypeError, ValueError):
                    pass
        # Disable Word auto-spacing that adds large gaps around body text.
        for attr in ("beforeAutospacing", "afterAutospacing"):
            key = qn(f"w:{attr}")
            if spacing.get(key) is not None:
                spacing.set(key, "0")


def _ensure_renderable_arabic_fonts(doc: Document, size_scale: float = 1.0) -> str:
    """
    Arabic Typesetting looks great in Word but is often missing on Linux.
    Remap document fonts to an installed Arabic font (Scheherazade New / Amiri)
    so LibreOffice preview/PDF conversion does not fall back to an ugly substitute.

    Scheherazade/Amiri glyphs are taller than Arabic Typesetting, so callers can
    pass size_scale < 1.0 to keep single-page layouts.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    font_name = _best_arabic_font()
    missing = {"Arabic Typesetting", "Andalus", "Sakkal Majalla", "Traditional Arabic"}

    def _patch_rfonts(rFonts) -> None:
        if rFonts is None:
            return
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            key = qn(f"w:{attr}")
            current = rFonts.get(key)
            if current is None or current in missing or current == "":
                rFonts.set(key, font_name)

    # Patch every style that still points at a Windows-only Arabic face.
    for style in doc.styles:
        try:
            rPr = style.element.get_or_add_rPr()
        except Exception:
            continue
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                rFonts.set(qn(f"w:{attr}"), font_name)
        else:
            _patch_rfonts(rFonts)
        try:
            style.font.name = font_name
        except Exception:
            pass

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name, size_scale=size_scale)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, font_name, size_scale=size_scale)

    if size_scale < 1.0:
        _tighten_paragraph_spacing(doc)

    return font_name


def generate_docx_certificate(
    template_id: str,
    data: dict[str, Any],
    output_path: Path | None = None,
) -> Path:
    template = get_template(template_id)
    if template.output_format != "docx":
        raise ValueError(f"Template {template_id} is not a DOCX template")

    resolved = _resolve_data(template, data)
    if template_id == "appreciation":
        # Template still has {{ year }}; derive it from the chosen dates.
        resolved["year"] = _derive_year(resolved)

    source = TEMPLATES_DIR / template.source_file
    if not source.exists():
        raise FileNotFoundError(f"Template file not found: {source}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if output_path is None:
        name_part = (
            resolved.get("recipient_name")
            or resolved.get("family_name")
            or resolved.get("council_location")
            or template.id
        )
        output_path = OUTPUT_DIR / f"{template.id}_{_safe_filename(name_part)}.docx"

    if _docx_uses_jinja(source):
        from docxtpl import DocxTemplate

        docx_tpl = DocxTemplate(str(source))
        docx_tpl.render(resolved)
        docx_tpl.save(str(output_path))
        # Re-open to normalize fonts / sizes for Linux/LibreOffice rendering.
        doc = Document(str(output_path))
        if template_id == "appreciation":
            _enlarge_appreciation_text(doc)
            # Only fix italic seal line — full Amiri remap made the cert 2 pages.
            _fix_arabic_pdf_fonts(doc, size_scale=0.9)
            _distribute_appreciation_layout(doc)
            doc.save(str(output_path))
        elif template_id in {"condolence_individual", "condolence_family", "congratulations"}:
            # Scheherazade is taller than Arabic Typesetting — shrink condolence
            # so the letter stays on a single page.
            scale = 0.65 if template_id.startswith("condolence") else 0.9
            _ensure_renderable_arabic_fonts(doc, size_scale=scale)
            if template_id.startswith("condolence"):
                # Slightly roomier lines after the compact font pass (not too much).
                _set_line_spacing(doc, line=276)
            if template_id == "congratulations":
                _enlarge_congratulations_fields(
                    doc,
                    resolved.get("recipient_title", ""),
                    resolved.get("council_location", ""),
                )
            doc.save(str(output_path))
        return output_path

    doc = Document(str(source))

    if template_id == "condolence_family":
        _apply_docx_rules(doc, DOCX_RULES["condolence_family"], resolved)
        _trim_condolence_doc(doc, "family")
    elif template_id == "condolence_individual":
        _apply_docx_rules(doc, DOCX_RULES["condolence_individual"], resolved)
        _trim_condolence_doc(doc, "individual")
    else:
        _apply_docx_rules(doc, DOCX_RULES[template_id], resolved)

    if template_id.startswith("condolence"):
        _ensure_renderable_arabic_fonts(doc, size_scale=0.65)
        _set_line_spacing(doc, line=276)
    elif template_id == "congratulations":
        _ensure_renderable_arabic_fonts(doc, size_scale=0.9)

    doc.save(str(output_path))
    return output_path


def _insert_arabic_text(
    page: fitz.Page,
    text: str,
    x: float,
    y: float,
    font_size: float = 14,
    bold: bool = False,
    max_width: float = 280,
) -> None:
    font_path = ARABIC_FONT_BOLD if bold else ARABIC_FONT
    display = _prepare_pdf_text(text)
    font = fitz.Font(fontfile=font_path)
    text_width = font.text_length(display, fontsize=font_size)
    start_x = max(x - max_width, x - text_width)
    page.insert_text(
        (start_x, y),
        display,
        fontname="dejavu",
        fontfile=font_path,
        fontsize=font_size,
        color=(0, 0, 0),
    )


def ensure_honorary_landscape_template() -> Path:
    """Return the landscape honorary PDF, building it from portrait if needed."""
    landscape = TEMPLATES_DIR / HONORARY_LANDSCAPE_SOURCE
    if landscape.exists():
        return landscape

    portrait = TEMPLATES_DIR / HONORARY_PORTRAIT_SOURCE
    if not portrait.exists():
        raise FileNotFoundError(
            f"Template file not found: {landscape} (and no portrait fallback at {portrait})"
        )

    src = fitz.open(str(portrait))
    page_w, page_h = src[0].rect.width, src[0].rect.height
    out = fitz.open()
    page = out.new_page(width=page_h, height=page_w)
    page.show_pdf_page(page.rect, src, 0, rotate=270)
    out.save(str(landscape))
    out.close()
    src.close()
    return landscape


# Landscape page 834×568, rotation=0 — x = right edge of each dotted line.
HONORARY_PDF_FIELDS = {
    "recipient_name": {"x": 600, "y": 359, "font_size": 14, "bold": True, "max_width": 360},
    "rank": {"x": 620, "y": 394, "font_size": 13, "bold": False, "max_width": 130},
    "job": {"x": 288, "y": 390, "font_size": 13, "bold": False, "max_width": 160},
    "issue_date": {"x": 600, "y": 496, "font_size": 12, "bold": False, "max_width": 120},
}


def generate_pdf_certificate(
    template_id: str,
    data: dict[str, Any],
    output_path: Path | None = None,
) -> Path:
    template = get_template(template_id)
    if template.output_format != "pdf":
        raise ValueError(f"Template {template_id} is not a PDF template")

    resolved = _resolve_data(template, data)
    source = ensure_honorary_landscape_template()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(str(source))
    page = doc[0]

    for field_id, placement in HONORARY_PDF_FIELDS.items():
        _insert_arabic_text(
            page,
            resolved[field_id],
            x=placement["x"],
            y=placement["y"],
            font_size=placement["font_size"],
            bold=placement["bold"],
            max_width=placement.get("max_width", 280),
        )

    if output_path is None:
        output_path = OUTPUT_DIR / f"{template.id}_{_safe_filename(resolved['recipient_name'])}.pdf"

    doc.save(str(output_path))
    doc.close()
    return output_path


def generate_certificate(
    template_id: str,
    data: dict[str, Any],
    output_path: Path | None = None,
) -> Path:
    template = get_template(template_id)
    if template.output_format == "docx":
        return generate_docx_certificate(template_id, data, output_path)
    if template.output_format == "pdf":
        return generate_pdf_certificate(template_id, data, output_path)
    raise ValueError(f"Unsupported output format: {template.output_format}")


def generate_from_csv(
    template_id: str,
    csv_path: str | Path,
    name_columns: tuple[str, str] | None = None,
    extra_fields: dict[str, str] | None = None,
) -> list[Path]:
    import csv

    template = get_template(template_id)
    csv_path = Path(csv_path)
    outputs: list[Path] = []

    with open(csv_path, newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            data = _map_row_to_data(template, row, extra_fields, name_columns)
            outputs.append(generate_certificate(template_id, data))
    return outputs


FIELD_COLUMN_ALIASES: dict[str, list[str]] = {
    "recipient_name": ["recipient_name", "nom", "name", "nom_complet", "nom complet", "full_name"],
    "family_name": ["family_name", "nom_famille", "famille"],
    "rank": ["rank", "grade", "rang", "الرتبة"],
    "job": ["job", "poste", "fonction", "position", "الوظيفة"],
    "issue_date": ["issue_date", "date", "date_emission", "تاريخ"],
    "rank_workplace": ["rank_workplace", "grade_lieu", "grade et lieu de travail"],
    "deceased_name": ["deceased_name", "nom_defunt", "defunt"],
    "training_field": ["training_field", "domaine", "domaine_formation"],
    "date_from": ["date_from", "date_debut", "debut"],
    "date_to": ["date_to", "date_fin", "fin"],
    "year": ["year", "annee", "année"],
    "coordinating_orgs": ["coordinating_orgs", "organismes", "organismes_coordinateurs"],
    "council_location": ["council_location", "conseil", "conseil_magistrature", "council location"],
    "recipient_title": ["recipient_title", "full_name", "full name", "name_title", "civilite", "civilité", "title"],
    "title": ["title", "civilite", "civilité"],
}


def _normalize_column_name(value: str) -> str:
    return re.sub(r"\s+", "_", value.strip().lower())


def _map_row_to_data(
    template: CertificateTemplate,
    row: dict[str, Any],
    extra_fields: dict[str, str] | None = None,
    name_columns: tuple[str, str] | None = None,
) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for key, value in row.items():
        if value is None:
            continue
        text = str(value).strip()
        if text:
            normalized[_normalize_column_name(str(key))] = text

    data: dict[str, str] = dict(extra_fields or {})

    if name_columns:
        first_col, last_col = name_columns
        first = normalized.get(_normalize_column_name(first_col), normalized.get("first_name", ""))
        last = normalized.get(_normalize_column_name(last_col), normalized.get("last_name", ""))
        if first or last:
            data.setdefault("recipient_name", f"{first} {last}".strip())

    alias_to_field: dict[str, str] = {}
    for field_id, aliases in FIELD_COLUMN_ALIASES.items():
        for alias in aliases:
            alias_to_field[_normalize_column_name(alias)] = field_id

    # Also accept the UI labels used as Excel column headers.
    for field in template.fields:
        for label in (field.id, field.label_fr, field.label_en, field.label_ar):
            if label:
                alias_to_field[_normalize_column_name(label)] = field.id

    for column, value in normalized.items():
        field_id = alias_to_field.get(column, column)
        if field_id in {f.id for f in template.fields}:
            data[field_id] = value

    return data


def read_spreadsheet_rows(path: str | Path) -> list[dict[str, str]]:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in {".xlsx", ".xlsm", ".xltx"}:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        sheet = workbook.active
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            return []
        headers = [str(cell or "").strip() for cell in rows[0]]
        records: list[dict[str, str]] = []
        for row in rows[1:]:
            if not row or not any(cell is not None and str(cell).strip() for cell in row):
                continue
            record: dict[str, str] = {}
            for header, cell in zip(headers, row):
                if not header or cell is None or (isinstance(cell, str) and not cell.strip()):
                    continue
                if hasattr(cell, "strftime"):
                    record[header] = cell.strftime("%d/%m/%Y")
                else:
                    text = str(cell).strip()
                    if text:
                        record[header] = text
            if record:
                records.append(record)
        workbook.close()
        return records

    if suffix == ".csv":
        import csv

        with open(path, newline="", encoding="utf-8-sig") as handle:
            return [dict(row) for row in csv.DictReader(handle) if any(v.strip() for v in row.values() if v)]

    raise ValueError(f"Format non pris en charge: {suffix}")


def generate_from_spreadsheet(
    template_id: str,
    spreadsheet_path: str | Path,
    extra_fields: dict[str, str] | None = None,
    name_columns: tuple[str, str] | None = None,
) -> list[Path]:
    template = get_template(template_id)
    rows = read_spreadsheet_rows(spreadsheet_path)
    if not rows:
        raise ValueError("Le fichier est vide ou ne contient aucune ligne de données.")

    outputs: list[Path] = []
    for index, row in enumerate(rows, start=1):
        try:
            data = _map_row_to_data(template, row, extra_fields, name_columns)
            outputs.append(generate_certificate(template_id, data))
        except Exception as exc:
            raise ValueError(f"Erreur à la ligne {index}: {exc}") from exc
    return outputs


def build_excel_template_bytes(template_id: str, *, with_examples: bool = True) -> bytes:
    """Build an Excel lot template with header + example rows (template and reference)."""
    from io import BytesIO

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    template = get_template(template_id)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Certificats"

    visible_fields = [
        field
        for field in template.fields
        if not (template.output_format == "pdf" and field.id == "title")
    ]
    headers = [field.label_fr or field.label_en for field in visible_fields]
    sheet.append(headers)

    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    examples: dict[str, list[list[str]]] = {
        "appreciation": [
            [
                "أحمد بن محمد",
                "التحول الرقمي",
                "01/03/2026",
                "15/03/2026",
                "وزارة العدل, المديرية العامة للموارد البشرية",
            ],
            [
                "فاطمة الزهراء",
                "القانون الإداري",
                "01/01/2026",
                "30/06/2026",
                "مجلس قضاء الجزائر",
            ],
        ],
        "condolence_individual": [
            ["إيمان بن كسيور", "قاضية - مجلس قضاء الجزائر", "أحمد بن محمد"],
            ["كريم بن سعيد", "مستشار - محكمة الجزائر", "محمد الأمين"],
        ],
        "congratulations": [
            ["أحمد بن محمد | السيد", "الجزائر العاصمة"],
            ["فاطمة الزهراء | السيدة", "وهران"],
        ],
        "honorary": [
            ["كريم بن سعيد", "مستشار", "رئيس قسم", date.today().strftime("%d/%m/%Y")],
            ["سارة بن علي", "قاضية", "رئيسة قسم", date.today().strftime("%d/%m/%Y")],
        ],
    }

    if with_examples:
        for row in examples.get(template_id, []):
            sheet.append(row)

    for index, field in enumerate(visible_fields, start=1):
        width = max(18, len(headers[index - 1]) + 4)
        if field.id in {"coordinating_orgs", "recipient_title", "rank_workplace"}:
            width = 42
        sheet.column_dimensions[get_column_letter(index)].width = width

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def write_bulk_excel_templates(output_dir: str | Path | None = None) -> list[Path]:
    """Write one Excel lot template per certificate type (headers + example rows)."""
    output_dir = Path(output_dir or BULK_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for template in list_templates():
        path = output_dir / f"modele_{template.id}.xlsx"
        path.write_bytes(build_excel_template_bytes(template.id, with_examples=True))
        written.append(path)
        # Remove legacy separate example files if present.
        legacy = output_dir / f"modele_{template.id}_exemple.xlsx"
        if legacy.exists():
            legacy.unlink()
    return written


def generate_preview_bytes(template_id: str, data: dict[str, Any]) -> bytes:
    output = generate_certificate(template_id, data)
    return output.read_bytes()


def pdf_to_png_bytes(pdf_path: str | Path, zoom: float = 1.3) -> bytes:
    doc = fitz.open(str(pdf_path))
    try:
        pixmap = doc[0].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        return pixmap.tobytes("png")
    finally:
        doc.close()


def _hidden_subprocess_kwargs() -> dict:
    """Kwargs so Windows does not flash CMD / console windows for child processes."""
    import os
    import subprocess

    if os.name != "nt":
        return {}

    startupinfo = subprocess.STARTUPINFO()
    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    startupinfo.wShowWindow = 0  # SW_HIDE
    # DETACHED_PROCESS (0x8) + CREATE_NO_WINDOW — hides soffice/taskkill consoles.
    creationflags = 0x00000008 | getattr(subprocess, "CREATE_NO_WINDOW", 0x08000000)
    return {
        "startupinfo": startupinfo,
        "creationflags": creationflags,
    }


def _run_hidden(cmd: list[str], **kwargs):
    """subprocess.run with console windows suppressed on Windows."""
    import subprocess

    for key, value in _hidden_subprocess_kwargs().items():
        kwargs.setdefault(key, value)
    return subprocess.run(cmd, **kwargs)


def _libreoffice_command() -> str:
    """Resolve LibreOffice binary on Linux/macOS/Windows."""
    import os
    import shutil

    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            # Prefer .exe over .com on Windows (.com is a console stub → CMD flash).
            if os.name == "nt" and found.lower().endswith(".com"):
                exe = Path(found).with_suffix(".exe")
                if exe.exists():
                    return str(exe)
            return found

    if os.name == "nt":
        candidates = [
            Path(os.environ.get("PROGRAMFILES", r"C:\Program Files"))
            / "LibreOffice"
            / "program"
            / "soffice.exe",
            Path(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"))
            / "LibreOffice"
            / "program"
            / "soffice.exe",
            Path(os.environ.get("LOCALAPPDATA", ""))
            / "Programs"
            / "LibreOffice"
            / "program"
            / "soffice.exe",
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)

    raise RuntimeError(
        "LibreOffice introuvable. Installez-le pour la conversion PDF "
        "(https://www.libreoffice.org/)."
    )


def _terminate_libreoffice_processes() -> None:
    """Best-effort cleanup of hung LibreOffice instances (common on Windows)."""
    import os

    if os.name == "nt":
        for image in ("soffice.exe", "soffice.bin", "oosplash.exe", "soffice.com"):
            _run_hidden(
                ["taskkill", "/F", "/IM", image, "/T"],
                capture_output=True,
                check=False,
            )
        return

    import subprocess

    subprocess.run(["pkill", "-9", "soffice"], capture_output=True, check=False)
    subprocess.run(["pkill", "-9", "libreoffice"], capture_output=True, check=False)


def docx_to_pdf_bytes(docx_path: str | Path) -> bytes:
    """Convert a DOCX to PDF bytes via LibreOffice.

    Optimizations:
    - Disk cache keyed by source mtime (instant on repeat)
    - Persistent LibreOffice user profile (avoids cold-start every call)
    - Light italic-Arabic patch only (no full font remap)
    - Hidden subprocesses on Windows (no CMD spam)
    """
    import os
    import shutil
    import tempfile
    import time

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    PREVIEW_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    st = docx_path.stat()
    cache_pdf = PREVIEW_CACHE_DIR / f"{docx_path.stem}_{int(st.st_mtime)}_{st.st_size}.pdf"
    if cache_pdf.exists() and cache_pdf.stat().st_size > 0:
        return cache_pdf.read_bytes()

    fonts_dir = BASE_DIR / "fonts"
    env = os.environ.copy()
    if fonts_dir.exists():
        existing = env.get("SAL_FONTPATH", "")
        env["SAL_FONTPATH"] = str(fonts_dir) if not existing else f"{fonts_dir}{os.pathsep}{existing}"
    env.setdefault("SAL_NO_NAG_DIALOGS", "1")

    soffice = _libreoffice_command()
    # Reuse one profile — creating a fresh profile each convert is the main slowdown.
    profile_root = WRITE_DIR / ".lo_user_profile"
    profile_root.mkdir(parents=True, exist_ok=True)
    profile_uri = profile_root.resolve().as_uri()

    with tempfile.TemporaryDirectory(prefix="cert_pdf_") as tmp:
        tmp_dir = Path(tmp)
        work_docx = tmp_dir / docx_path.name
        shutil.copy2(docx_path, work_docx)
        # Cheap tofu fix: only italic Arabic runs (body Arial stays as-is).
        try:
            doc = Document(str(work_docx))
            _fix_arabic_pdf_fonts(doc, size_scale=0.9)
            doc.save(str(work_docx))
        except Exception:
            pass

        cmd = [
            soffice,
            "--headless",
            "--invisible",
            "--nologo",
            "--norestore",
            "--nofirststartwizard",
            "--nolockcheck",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(tmp_dir),
            str(work_docx.resolve()),
        ]

        import subprocess

        run_kwargs: dict = {
            "capture_output": True,
            "text": True,
            "timeout": 120,
            "check": False,
            "env": env,
            "cwd": str(tmp_dir),
        }

        last_error = ""
        for attempt in range(2):
            try:
                result = _run_hidden(cmd, **run_kwargs)
            except subprocess.TimeoutExpired:
                _terminate_libreoffice_processes()
                last_error = "LibreOffice a dépassé le délai (120 s)"
                if attempt == 0:
                    time.sleep(0.4)
                    continue
                raise RuntimeError(
                    "Conversion DOCX → PDF échouée : délai dépassé. "
                    "Fermez LibreOffice s'il est ouvert, puis réessayez."
                ) from None

            pdf_path = tmp_dir / f"{work_docx.stem}.pdf"
            if result.returncode == 0 and pdf_path.exists():
                pdf_bytes = pdf_path.read_bytes()
                try:
                    cache_pdf.write_bytes(pdf_bytes)
                except OSError:
                    pass
                return pdf_bytes

            last_error = (result.stderr or result.stdout or "").strip()
            if attempt == 0:
                _terminate_libreoffice_processes()
                time.sleep(0.4)
                continue

        raise RuntimeError(
            f"Conversion DOCX → PDF échouée: {last_error or 'LibreOffice error'}"
        )


def pdf_bytes_to_png_bytes(pdf_bytes: bytes, zoom: float = 1.3) -> bytes:
    """Render the first page of PDF bytes to PNG."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        pixmap = doc[0].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        return pixmap.tobytes("png")
    finally:
        doc.close()


def docx_to_png_bytes(docx_path: str | Path, zoom: float = 1.3) -> bytes:
    """Convert a DOCX to PNG via LibreOffice (DOCX → PDF → PNG)."""
    return pdf_bytes_to_png_bytes(docx_to_pdf_bytes(docx_path), zoom=zoom)


def document_to_png_bytes(path: str | Path, zoom: float = 1.3) -> bytes:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return pdf_to_png_bytes(path, zoom=zoom)
    if suffix == ".docx":
        return docx_to_png_bytes(path, zoom=zoom)
    raise ValueError(f"Aperçu non pris en charge pour: {suffix}")


def static_template_preview_path(template_id: str) -> Path | None:
    """Bundled PNG preview for official templates (instant, no LibreOffice)."""
    candidate = STATIC_PREVIEWS_DIR / f"{template_id}.png"
    return candidate if candidate.exists() else None


def cached_document_preview(path: str | Path, zoom: float = 1.1) -> bytes:
    """Disk-cached PNG preview; first DOCX conversion may still need LibreOffice."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    PREVIEW_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    mtime = int(path.stat().st_mtime)
    zoom_key = str(zoom).replace(".", "_")
    cache_path = PREVIEW_CACHE_DIR / f"{path.stem}_{mtime}_{zoom_key}.png"
    if cache_path.exists():
        return cache_path.read_bytes()

    png = document_to_png_bytes(path, zoom=zoom)
    cache_path.write_bytes(png)
    return png
